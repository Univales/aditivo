import json, os
from num2words import num2words
from datetime import datetime, timedelta
from datetime import date
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches

username = os.getlogin()
caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
altera_vencimento = 'altera_vencimento.docx'

print('teste python')
exit()
#Cria o documento
doc = Document()

#Variaveis Emitente
strTitulo = 'C20234510-2'
strEmitente = 'WELINTON NASCIMENTO SANTOS CRUZ'.upper()
strNacioEmitente = 'brasileira'.upper()
strEstCivilEmitente = 'casado'.upper()
strDataDocumento = '28 de julho de 2022'
strPaiEmitente='JONES DA SILVA CRUZ'.upper()
strMaeEmitente='EDNA DO NASCIMENTO SANTOS'.upper()
strProfEmitente='PRODUTOR AGROPECUÁRIO'
strLogEmitente='EST. PA JURUENA'
strNumCasaEmitente='0'
strBairroEmitente='RURAL'
strMunEmitent='COTRIGUACU'
strUFEmitente='MT'
strCEPEmitente='78330-000'
strCPFEmitente='022.362.241-92'
strRGEmitente='000896344'
strOrgaoRGEmitente = 'SESDEC'
strUFRG = 'RO'
strTelEmitente='(66) 99438-0493'
strEmailEmitente='borrachacabral@gmail.com'
strAditivoa='Cédula de Crédito Bancário'.upper()
#variaveis caso for alteração de taxa
taxa_ano = 16.56
taxa_mes = taxa_ano/12
#variaveis caso for troca de garantia
garantia_retirada = 'IMOVEL RURAL LOTE 09A, GLEBA 05, DENOMINADO RANCHO TORTOLA FERREIRA, COM AREA DE 17,1870HA, NO MUNICIPIO DE SÃO MIGUEL DO GUAPORÉ-RO. REGISTRADO SOB MATRÍCULA DE N° 10.419, NO LIVRO 2-REGISTRO GERAL, NO REGISTRO DE IMOVEIS DA COMARCA DE SÃO MIGUEL DO GUAPORÉ/RO. COM SEUS LIMITES E CONFRONTAÇÕES CONFORME MATRÍCULA EM ANEXO. ADQUIRIDO SOB ESCRITURA PUBLICA DE COMPRA E VENDA LAVRADA NAS FOLHAS 024/025, DO LIVRO 57-E, NO TABELIONATO DE NOTAS DE SÃO MIGUEL DO GUAPORÉ-RO, EM 03/07/2023, NO REGISTRO DE IMOVEIS, TITULOS E DOCUMENTOS, CIVIL DAS PESSOAS JURIDICAS E TABELIONATO DE PROTESTO COMARCA DE SÃO MIGUEL DO GUAPORÉ/RO. IMOVEL AVALIADO CONFORME LIQUIDEZ FORCADA PARA EFEITO DE VENDA EM PUBLICO LEILAO NO VALOR DE R$ 909.056,00 (NOVECENTOS E NOVE MIL E CINQUENTA E SEIS REAIS).'
nova_garantia ='IMOVEL RURAL LOTE 09A, GLEBA 05, DENOMINADO RANCHO TORTOLA FERREIRA, COM AREA DE 17,1870HA, NO MUNICIPIO DE SÃO MIGUEL DO GUAPORÉ-RO. REGISTRADO SOB MATRÍCULA DE N° 10.419, NO LIVRO 2-REGISTRO GERAL, NO REGISTRO DE IMOVEIS DA COMARCA DE SÃO MIGUEL DO GUAPORÉ/RO. COM SEUS LIMITES E CONFRONTAÇÕES CONFORME MATRÍCULA EM ANEXO. ADQUIRIDO SOB ESCRITURA PUBLICA DE COMPRA E VENDA LAVRADA NAS FOLHAS 024/025, DO LIVRO 57-E, NO TABELIONATO DE NOTAS DE SÃO MIGUEL DO GUAPORÉ-RO, EM 03/07/2023, NO REGISTRO DE IMOVEIS, TITULOS E DOCUMENTOS, CIVIL DAS PESSOAS JURIDICAS E TABELIONATO DE PROTESTO COMARCA DE SÃO MIGUEL DO GUAPORÉ/RO. IMOVEL AVALIADO CONFORME LIQUIDEZ FORCADA PARA EFEITO DE VENDA EM PUBLICO LEILAO NO VALOR DE R$ 909.056,00 (NOVECENTOS E NOVE MIL E CINQUENTA E SEIS REAIS).'
local_deposito=strMunEmitent
#variaveis caso for troca de vinculo
strdataemissao = '23/09/2023'
strValorParcela = '153.200,76'
strQtdParcelas = '1'
strdataparcela = '23/09/2023'
strNovoTitulo = 'C30433700-1'
strTaxaMes = '2,300000'
strTaxaAno = '31,373450'

#variaveis caso a garantia for imovel
strCarencia = '60'
strValorImovel = '909.056,00'
strTituloExcGarantia = 'C122234517'
strValorExcGarantia = '500.000,00'
strVencFinal = '22/12/2025'
strTxJurosExc = '0,70'
strInstrumentoExc = 'CÉDULA DE CRÉDITO BANCÁRIO'.upper()
#Variaveis Avalista
strjson_avalista = '[{"vConta":"367781","vNome":"VINICIUS PARREIRA BATISTA FERREIRA","vCPF":"05795880175","vRG":"26111136","vProfissao":"Assistente administrativo","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"SATELITE","vNumero":"130A","vBairro":"MODULO 4","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"GERALDA JOANA DARK PARREIRA","vPai":"LUIZ BATISTA FERREIRA","vUF":"MT","vCep":"98320000","vOrgao":"SEJUSP","vEmail":"parreira2000@hotmail.com","vUFdoc":"MT","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"060606","vNome":"CAYO CESAR CAVALCANTE GARCES","vCPF":"04407638192","vRG":"05529267512","vProfissao":"Analista de negócios","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"BOA VISTA","vEndereco":"MATO GROSSO","vNumero":"1923","vBairro":"MODULO 05","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"MARCIA CAVALCANTE","vPai":"WILBER TAPIA GARCES ","vUF":"MT","vCep":"78320000","vOrgao":"DETRAN","vEmail":"CAYOCAVALCANTE@HOTMAIL.COM","vUFdoc":"RR","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"432047","vNome":"RAPHAEL EGNALDO LEANDRO","vCPF":"05464116141","vRG":"23066032","vProfissao":"Inspetor de qualidade","vEstCivil":"CASADO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"PRESIDENTE PRUDENTE","vNumero":"106W","vBairro":"MODULO 06","vNomeConjuge":"CAROLINA MARCHEZAN AVELINO","vCPFConjuge":"04665118175","vMae":"ERNESTA DA SILVA ARAUJO","vPai":"WAGNER LEANDRO","vUF":"MT","vCep":"78320000","vOrgao":"SSP","vEmail":"RELEANDRO15@GMAIL.COM","vUFdoc":"MT","vRegCas":"COMUNHÃO PARCIAL DE BENS","vRGConj":"21981795","vEmissorConj":"SSP","vUFDocConj":"MT","vMaeConj":null,"vPaiConj":null,"vNacConj":"BRASILEIRA","vProfConj":"Gerente de produtos bancários","vCEPConj":"78320000","vEmailConj":null,"vEndConj":"DR ULISSES GUIMARAES","vNumConj":"1920 W","vBairroConj":"MODULO 06","vCidConj":"JUINA","vUFConj":"MT"}]'

strjson_excavalista = '[{"vConta":"367781","vNome":"VINICIUS PARREIRA BATISTA FERREIRA","vCPF":"05795880175","vRG":"26111136","vProfissao":"Assistente administrativo","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"SATELITE","vNumero":"130A","vBairro":"MODULO 4","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"GERALDA JOANA DARK PARREIRA","vPai":"LUIZ BATISTA FERREIRA","vUF":"MT","vCep":"98320000","vOrgao":"SEJUSP","vEmail":"parreira2000@hotmail.com","vUFdoc":"MT","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"060606","vNome":"CAYO CESAR CAVALCANTE GARCES","vCPF":"04407638192","vRG":"05529267512","vProfissao":"Analista de negócios","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"BOA VISTA","vEndereco":"MATO GROSSO","vNumero":"1923","vBairro":"MODULO 05","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"MARCIA CAVALCANTE","vPai":"WILBER TAPIA GARCES ","vUF":"MT","vCep":"78320000","vOrgao":"DETRAN","vEmail":"CAYOCAVALCANTE@HOTMAIL.COM","vUFdoc":"RR","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"432047","vNome":"RAPHAEL EGNALDO LEANDRO","vCPF":"05464116141","vRG":"23066032","vProfissao":"Inspetor de qualidade","vEstCivil":"CASADO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"PRESIDENTE PRUDENTE","vNumero":"106W","vBairro":"MODULO 06","vNomeConjuge":"CAROLINA MARCHEZAN AVELINO","vCPFConjuge":"04665118175","vMae":"ERNESTA DA SILVA ARAUJO","vPai":"WAGNER LEANDRO","vUF":"MT","vCep":"78320000","vOrgao":"SSP","vEmail":"RELEANDRO15@GMAIL.COM","vUFdoc":"MT","vRegCas":"COMUNHÃO PARCIAL DE BENS","vRGConj":"21981795","vEmissorConj":"SSP","vUFDocConj":"MT","vMaeConj":null,"vPaiConj":null,"vNacConj":"BRASILEIRA","vProfConj":"Gerente de produtos bancários","vCEPConj":"78320000","vEmailConj":null,"vEndConj":"DR ULISSES GUIMARAES","vNumConj":"1920 W","vBairroConj":"MODULO 06","vCidConj":"JUINA","vUFConj":"MT"}]'

strjson_incavalista = '[{"vConta":"367781","vNome":"VINICIUS PARREIRA BATISTA FERREIRA","vCPF":"05795880175","vRG":"26111136","vProfissao":"Assistente administrativo","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"SATELITE","vNumero":"130A","vBairro":"MODULO 4","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"GERALDA JOANA DARK PARREIRA","vPai":"LUIZ BATISTA FERREIRA","vUF":"MT","vCep":"98320000","vOrgao":"SEJUSP","vEmail":"parreira2000@hotmail.com","vUFdoc":"MT","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"060606","vNome":"CAYO CESAR CAVALCANTE GARCES","vCPF":"04407638192","vRG":"05529267512","vProfissao":"Analista de negócios","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"BOA VISTA","vEndereco":"MATO GROSSO","vNumero":"1923","vBairro":"MODULO 05","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"MARCIA CAVALCANTE","vPai":"WILBER TAPIA GARCES ","vUF":"MT","vCep":"78320000","vOrgao":"DETRAN","vEmail":"CAYOCAVALCANTE@HOTMAIL.COM","vUFdoc":"RR","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"432047","vNome":"RAPHAEL EGNALDO LEANDRO","vCPF":"05464116141","vRG":"23066032","vProfissao":"Inspetor de qualidade","vEstCivil":"CASADO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"PRESIDENTE PRUDENTE","vNumero":"106W","vBairro":"MODULO 06","vNomeConjuge":"CAROLINA MARCHEZAN AVELINO","vCPFConjuge":"04665118175","vMae":"ERNESTA DA SILVA ARAUJO","vPai":"WAGNER LEANDRO","vUF":"MT","vCep":"78320000","vOrgao":"SSP","vEmail":"RELEANDRO15@GMAIL.COM","vUFdoc":"MT","vRegCas":"COMUNHÃO PARCIAL DE BENS","vRGConj":"21981795","vEmissorConj":"SSP","vUFDocConj":"MT","vMaeConj":null,"vPaiConj":null,"vNacConj":"BRASILEIRA","vProfConj":"Gerente de produtos bancários","vCEPConj":"78320000","vEmailConj":null,"vEndConj":"DR ULISSES GUIMARAES","vNumConj":"1920 W","vBairroConj":"MODULO 06","vCidConj":"JUINA","vUFConj":"MT"}]'

strjson_intergarantidor = '[{"vConta":"367781","vNome":"VINICIUS PARREIRA BATISTA FERREIRA","vCPF":"05795880175","vRG":"26111136","vProfissao":"Assistente administrativo","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"SATELITE","vNumero":"130A","vBairro":"MODULO 4","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"GERALDA JOANA DARK PARREIRA","vPai":"LUIZ BATISTA FERREIRA","vUF":"MT","vCep":"98320000","vOrgao":"SEJUSP","vEmail":"parreira2000@hotmail.com","vUFdoc":"MT","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"060606","vNome":"CAYO CESAR CAVALCANTE GARCES","vCPF":"04407638192","vRG":"05529267512","vProfissao":"Analista de negócios","vEstCivil":"SOLTEIRO","vNacionalidade":"BRASILEIRA","vCidade":"BOA VISTA","vEndereco":"MATO GROSSO","vNumero":"1923","vBairro":"MODULO 05","vNomeConjuge":null,"vCPFConjuge":null,"vMae":"MARCIA CAVALCANTE","vPai":"WILBER TAPIA GARCES ","vUF":"MT","vCep":"78320000","vOrgao":"DETRAN","vEmail":"CAYOCAVALCANTE@HOTMAIL.COM","vUFdoc":"RR","vRegCas":null,"vRGConj":null,"vEmissorConj":null,"vUFDocConj":null,"vMaeConj":null,"vPaiConj":null,"vNacConj":null,"vProfConj":null,"vCEPConj":null,"vEmailConj":null,"vEndConj":null,"vNumConj":null,"vBairroConj":null,"vCidConj":null,"vUFConj":null},{"vConta":"432047","vNome":"RAPHAEL EGNALDO LEANDRO","vCPF":"05464116141","vRG":"23066032","vProfissao":"Inspetor de qualidade","vEstCivil":"CASADO","vNacionalidade":"BRASILEIRA","vCidade":"JUINA","vEndereco":"PRESIDENTE PRUDENTE","vNumero":"106W","vBairro":"MODULO 06","vNomeConjuge":"CAROLINA MARCHEZAN AVELINO","vCPFConjuge":"04665118175","vMae":"ERNESTA DA SILVA ARAUJO","vPai":"WAGNER LEANDRO","vUF":"MT","vCep":"78320000","vOrgao":"SSP","vEmail":"RELEANDRO15@GMAIL.COM","vUFdoc":"MT","vRegCas":"COMUNHÃO PARCIAL DE BENS","vRGConj":"21981795","vEmissorConj":"SSP","vUFDocConj":"MT","vMaeConj":null,"vPaiConj":null,"vNacConj":"BRASILEIRA","vProfConj":"Gerente de produtos bancários","vCEPConj":"78320000","vEmailConj":null,"vEndConj":"DR ULISSES GUIMARAES","vNumConj":"1920 W","vBairroConj":"MODULO 06","vCidConj":"JUINA","vUFConj":"MT"}]'

#numero clausula
intClausula=0


#definindo o tipo da garantia
strNovaGarantia = 'imóvel'
strGarantiaRetirada = 'imóvel'

#meses
meses = {
    1: "Janeiro",
    2: "Fevereiro",
    3: "Março",
    4: "Abril",
    5: "Maio",
    6: "Junho",
    7: "Julho",
    8: "Agosto",
    9: "Setembro",
    10: "Outubro",
    11: "Novembro",
    12: "Dezembro"
}
#Parcelas altera vencimento
primParcela = 13
dataPrimParcela='30/05/2023'
dtPrimParcela=datetime.strptime(dataPrimParcela, "%d/%m/%Y")
ultParcela = 20
qtdParcelas = (ultParcela-primParcela)
periodicidade = 6


#==============================================================================================================================
if strjson_avalista:
    # Carrega a string JSON como uma lista de dicionários
    json_avalista = json.loads(strjson_avalista)
    #variaveis para incluir avalista
    strNomeAvalista = json_avalista[0]['vNome']
    strNacioAvalista = json_avalista[0]['vNacionalidade']
    strProfAvalista = json_avalista[0]['vProfissao']
    if json_avalista[0]['vRegCas'] != None:
        strEstCivilAvalista = json_avalista[0]['vEstCivil']+', sob '+ json_avalista[0]['vRegCas']
    strRGAvalista = json_avalista[0]['vRG']
    strCPFAvalista = json_avalista[0]['vCPF']
    strEndAvalista = json_avalista[0]['vEndereco']+', número '+json_avalista[0]['vNumero']+', '+json_avalista[0]['vBairro']+', '+json_avalista[0]['vCep']+', '+json_avalista[0]['vCidade']+', '+json_avalista[0]['vUF']

    # Conta o número de itens na lista
    qtdAvalistas = len(json_avalista)
    print('O aditivo possui '+str(qtdAvalistas)+' avalistas')

if strjson_excavalista:
# Carrega a string JSON como uma lista de dicionários
    json_excavalista = json.loads(strjson_excavalista)

    #variaveis para excluir avalista
    strNomeExcAvalista = json_excavalista[0]['vNome']
    strNacioExcAvalista = json_excavalista[0]['vNacionalidade']
    strProfExcAvalista = json_excavalista[0]['vProfissao']
    if json_excavalista[0]['vRegCas'] != None:
        strEstCivilExcAvalista = json_excavalista[0]['vEstCivil']+', sob '+ json_excavalista[0]['vRegCas']
    strRGExcAvalista = json_excavalista[0]['vRG']
    strCPFExcAvalista = json_excavalista[0]['vCPF']
    strEndExcAvalista = json_excavalista[0]['vEndereco']+', número '+json_excavalista[0]['vNumero']+', '+json_excavalista[0]['vBairro']+', '+json_excavalista[0]['vCep']+', '+json_excavalista[0]['vCidade']+', '+json_excavalista[0]['vUF']
    # Conta o número de itens na lista
    qtdExcAvalistas = len(json_excavalista)
    print('O aditivo possui '+str(qtdExcAvalistas)+' avalistas para excluir')

if strjson_incavalista:
# Carrega a string JSON como uma lista de dicionários
    json_incavalista = json.loads(strjson_incavalista)

    #variaveis para incluir avalista
    strNomeIncAvalista = json_incavalista[0]['vNome']
    strNacioIncAvalista = json_incavalista[0]['vNacionalidade']
    strProfIncAvalista = json_incavalista[0]['vProfissao']
    if json_incavalista[0]['vRegCas'] != None:
        strEstCivilIncAvalista = json_incavalista[0]['vEstCivil']+', sob '+ json_incavalista[0]['vRegCas']
    strRGIncAvalista = json_incavalista[0]['vRG']
    strCPFIncAvalista = json_incavalista[0]['vCPF']
    strEndIncAvalista = json_incavalista[0]['vEndereco']+', número '+json_incavalista[0]['vNumero']+', '+json_incavalista[0]['vBairro']+', '+json_incavalista[0]['vCep']+', '+json_incavalista[0]['vCidade']+', '+json_incavalista[0]['vUF']
    # Conta o número de itens na lista
    qtdIncAvalistas = len(json_incavalista)
    print('O aditivo possui '+str(qtdIncAvalistas)+' avalistas para incluir')

if strjson_intergarantidor:
    # Carrega a string JSON como uma lista de dicionários
    json_intergarantidor = json.loads(strjson_intergarantidor)

    # Conta o número de itens na lista
    qtdIntergarantidor = len(json_intergarantidor)
    print('O aditivo possui '+str(qtdIntergarantidor)+' Intervenientes Garantidores')

#=================================================INICIO DO PROGRAMA============================================================================================================
#CABEÇALHO
txtCabecalho = 'Aditivo à '+strAditivoa+' Nº '+strTitulo+', emitida por '+strEmitente+' em favor da COOPERATIVA DE CRÉDITO, POUPANÇA E INVESTIMENTO UNIVALES - SICREDI UNIVALES MT/RO, CNPJ 70.431.630/0001-04, em '+strDataDocumento+'.'

#CREDOR
txtCredor = 'CREDOR: COOPERATIVA DE CRÉDITO POUPANÇA E INVESTIMENTO UNIVALES SICREDI UNIVALES MT/RO, instituição financeira brasileira, estabelecida no(a) AV. MATO GROSSO, 690N, município de JUINA-MT, inscrita no CNPJ sob número 70.431.630/0001-04.'

#EMITENTE
txtEmitente ='EMITENTE(S): '+strEmitente+', Nacionalidade '+strNacioEmitente+', '+strEstCivilEmitente+', filho(a) de '+strPaiEmitente+' e '+strMaeEmitente+', '+strProfEmitente+', residente e domiciliado(a) no(a) '+strLogEmitente+', '+strNumCasaEmitente+', bairro '+strBairroEmitente+', município de '+strMunEmitent+'-'+strUFEmitente+', '+strCEPEmitente+', inscrito no CPF sob n. '+strCPFEmitente+' e RG '+strRGEmitente+' - '+strOrgaoRGEmitente+'/'+strUFRG+', telefone '+strTelEmitente+', endereço eletrônico '+strEmailEmitente+'.'

#Assinatura Credor
txtAssCredor = "CREDOR: COOPERATIVA DE CRÉDITO, POUPANÇA E INVESTIMENTO UNIVALES - SICREDI UNIVALES MT/RO\n CNPJ.: 70.431.630/0001-04\n\n\n"

#Assinatura Emitente
txtAssEmitente = "EMITENTE(S)/ASSOCIADO(S)\n\n\n\n\n\nNome: "+strEmitente+"\nCPF: "+strCPFEmitente

#Rodapé
hoje = date.today()
data_formatada = hoje.strftime("%d/%m/%Y")
txtRodape=(strMunEmitent+'-'+strUFEmitente+', '+ str(hoje.day)+' de '+meses[hoje.month]+ ' de '+str(hoje.year))
#====================================DEFINE OS ITENS DAS GRIDS======================================================================
#Definir os avalistas
def add_avalistas(doc, json_avalista):
    pAvalista = doc.add_paragraph('Avalista(s): ')
    for I in json_avalista:
        texto = f"{I['vNome']}, Nacionalidade {I['vNacionalidade']}, "
        if I['vEstCivil'].lower().startswith('casad'):
            texto+=f"{I['vEstCivil']}, pelo regime de {I['vRegCas']}, "
        else:
            texto+=f"{I['vEstCivil']}, "
            
        texto+=f"filho(a) de {I['vPai']} e {I['vMae']}, {I['vProfissao']}, residente e domiciliado(a) no(a) {I['vEndereco']}, {I['vNumero']}, bairro {I['vBairro']}, município de {I['vCidade']} - {I['vUF']}, {I['vCep']}, CPF {I['vCPF']} e RG {I['vRG']} - {I['vOrgao']}/{I['vUFdoc']}, endereço eletrônico {I['vEmail']}."
        doc.add_paragraph(texto).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #verifica se o regime de casamento é comunhão parcial de bens ou comunhão universal de bens, pois nesses casos é necessario coletar assinatura do conjuge
        if I['vRegCas'] is not None and I['vRegCas'].lower() in ['comunhão parcial de bens', 'comunhão universal de bens']:
            pAvalista = doc.add_paragraph('Cônjuge: ')
            texto = f"{I['vNomeConjuge']}, CPF {I['vCPFConjuge']} e RG {I['vRGConj']} Nacionalidade {I['vNacConj']}, "
            if I['vEstCivil'].lower().startswith('casad'):
                texto+=f"{I['vEstCivil']}, pelo regime de {I['vRegCas']}, "
            else:
                texto+=f"{I['vEstCivil']}, "
            texto+=f"filho(a) de {I['vPaiConj']} e {I['vMaeConj']}, {I['vProfConj']}, residente e domiciliado(a) no(a) {I['vEndConj']}, {I['vNumConj']}, bairro {I['vBairroConj']}, município de {I['vCidConj']} - {I['vUFConj']}, {I['vCEPConj']}, CPF {I['vCPFConjuge']} e RG {I['vRGConj']} - {I['vEmissorConj']}/{I['vUFDocConj']}, endereço eletrônico {I['vEmailConj']}."
            doc.add_paragraph(texto).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

#Definir os avalistas a excluir
def excluir_avalistas(doc, json_excavalista):
    pAvalista = doc.add_paragraph('Avalista(s): ')
    for I in json_excavalista:
        texto = f"{I['vNome']}, Nacionalidade {I['vNacionalidade']}, "
        if I['vEstCivil'].lower().startswith('casad'):
            texto+=f"{I['vEstCivil']}, pelo regime de {I['vRegCas']}, "
        else:
            texto+=f"{I['vEstCivil']}, "
            
        texto+=f"filho(a) de {I['vPai']} e {I['vMae']}, {I['vProfissao']}, residente e domiciliado(a) no(a) {I['vEndereco']}, {I['vNumero']}, bairro {I['vBairro']}, município de {I['vCidade']} - {I['vUF']}, {I['vCep']}, CPF {I['vCPF']} e RG {I['vRG']} - {I['vOrgao']}/{I['vUFdoc']}, endereço eletrônico {I['vEmail']}."
        doc.add_paragraph(texto).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #verifica se o regime de casamento é comunhão parcial de bens ou comunhão universal de bens, pois nesses casos é necessario coletar assinatura do conjuge
        if I['vRegCas'] is not None and I['vRegCas'].lower() in ['comunhão parcial de bens', 'comunhão universal de bens']:
            pAvalista = doc.add_paragraph('Cônjuge: ')
            texto = f"{I['vNomeConjuge']}, CPF {I['vCPFConjuge']} e RG {I['vRGConj']} Nacionalidade {I['vNacConj']}, "
            if I['vEstCivil'].lower().startswith('casad'):
                texto+=f"{I['vEstCivil']}, pelo regime de {I['vRegCas']}, "
            else:
                texto+=f"{I['vEstCivil']}, "
            texto+=f"filho(a) de {I['vPaiConj']} e {I['vMaeConj']}, {I['vProfConj']}, residente e domiciliado(a) no(a) {I['vEndConj']}, {I['vNumConj']}, bairro {I['vBairroConj']}, município de {I['vCidConj']} - {I['vUFConj']}, {I['vCEPConj']}, CPF {I['vCPFConjuge']} e RG {I['vRGConj']} - {I['vEmissorConj']}/{I['vUFDocConj']}, endereço eletrônico {I['vEmailConj']}."
            doc.add_paragraph(texto).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY            

#Definir os intervenientes garantidores
def add_interveniente_garantidor(doc, json_intergarantidor):
    pAvalista = doc.add_paragraph('Devedor Fiduciante: ')
    for I in json_intergarantidor:
        texto = f"{I['vNome']}, Nacionalidade {I['vNacionalidade']}, "
        if I['vEstCivil'].lower().startswith('casad'):
            texto+=f"{I['vEstCivil']}, pelo regime de {I['vRegCas']}, "
        else:
            texto+=f"{I['vEstCivil']}, "
            
        texto+=f"filho(a) de {I['vPai']} e {I['vMae']}, {I['vProfissao']}, residente e domiciliado(a) no(a) {I['vEndereco']}, {I['vNumero']}, bairro {I['vBairro']}, município de {I['vCidade']} - {I['vUF']}, {I['vCep']}, CPF {I['vCPF']} e RG {I['vRG']} - {I['vOrgao']}/{I['vUFdoc']}, endereço eletrônico {I['vEmail']}."
        doc.add_paragraph(texto).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #verifica se o regime de casamento é comunhão parcial de bens ou comunhão universal de bens, pois nesses casos é necessario coletar assinatura do conjuge
        if I['vRegCas'] is not None and I['vRegCas'].lower() in ['comunhão parcial de bens', 'comunhão universal de bens']:
            pAvalista = doc.add_paragraph('Cônjuge: ')
            texto = f"{I['vNomeConjuge']}, CPF {I['vCPFConjuge']} e RG {I['vRGConj']} Nacionalidade {I['vNacConj']}, "
            if I['vEstCivil'].lower().startswith('casad'):
                texto+=f"{I['vEstCivil']}, pelo regime de {I['vRegCas']}, "
            else:
                texto+=f"{I['vEstCivil']}, "
            texto+=f"filho(a) de {I['vPaiConj']} e {I['vMaeConj']}, {I['vProfConj']}, residente e domiciliado(a) no(a) {I['vEndConj']}, {I['vNumConj']}, bairro {I['vBairroConj']}, município de {I['vCidConj']} - {I['vUFConj']}, {I['vCEPConj']}, CPF {I['vCPFConjuge']} e RG {I['vRGConj']} - {I['vEmissorConj']}/{I['vUFDocConj']}, endereço eletrônico {I['vEmailConj']}."
            doc.add_paragraph(texto).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  

#Definir Assinatura dos avalistas
def add_ass_avalistas(doc, json_avalista):
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run('\n\nPor aval ao(s) emitente(s)')
    for I in json_avalista:
        txtAssAval = f"\n\n\n\n\n\nNome: {I['vNome']}\nCPF: {I['vCPF']}"
        if I['vRegCas'] is not None and I['vRegCas'].lower() in ['comunhão parcial de bens', 'comunhão universal de bens']:
            txtAssAval+=f"\n\n\n\n\n\nNome: {I['vNomeConjuge']}\nCPF: {I['vCPFConjuge']}"
        pAssinaturas.add_run(txtAssAval)

#Definir Assinatura da Exclusão dos avalistas
def add_ass_exc_avalistas(doc, json_excavalista):
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run('\n\nPor aval ao(s) emitente(s)')
    for I in json_excavalista:
        txtAssAval = f"\n\n\n\n\n\nNome: {I['vNome']}\nCPF: {I['vCPF']}"
        if I['vRegCas'] is not None and I['vRegCas'].lower() in ['comunhão parcial de bens', 'comunhão universal de bens']:
            txtAssAval+=f"\n\n\n\n\n\nNome: {I['vNomeConjuge']}\nCPF: {I['vCPFConjuge']}"
        pAssinaturas.add_run(txtAssAval)

def add_ass_inter_garantidor(doc, json_intergarantidor):
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run('\n\nPor aval ao(s) emitente(s)')
    for I in json_intergarantidor:
        txtAssInter = f"\n\n\n\n\n\nNome: {I['vNome']}\nCPF: {I['vCPF']}"
        if I['vRegCas'] is not None and I['vRegCas'].lower() in ['comunhão parcial de bens', 'comunhão universal de bens']:
            txtAssInter+=f"\n\n\n\n\n\nNome: {I['vNomeConjuge']}\nCPF: {I['vCPFConjuge']}"
        pAssinaturas.add_run(txtAssInter)

#Paragrafos inclusivos: Esses parágrafos serão usados para incluir caso seja incluso mais de um tipo de aditivo==================================================================================================================================
def add_prorroga_data():
    pClausulasVeiculo = doc.add_paragraph()
    pClausulasVeiculo.add_run('3 - FORMA DE PAGAMENTO: As Partes resolvem alterar a data de vencimento da parcela 002, passando a ser 01/01/2024 parcela essa que incluem o principal e os encargos contratados.')
    pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
#===================================================================================================================================
#adiciona numeros nas páginas
def add_num_pag(doc):
    # Acesse a seção do documento
    section = doc.sections[0]

    # Crie um rodapé
    footer = section.footer

    # Adicione um campo de número de página e total de páginas ao rodapé
    field_code_page = 'PAGE'
    field_code_num_pages = 'NUMPAGES'
    docx_field_page = f'<w:fldSimple {nsdecls("w")} w:instr="{field_code_page}"/>'
    docx_field_num_pages = f'<w:fldSimple {nsdecls("w")} w:instr="{field_code_num_pages}"/>'
    footer.paragraphs[0]._p.append(parse_xml(docx_field_page))
    footer.paragraphs[0].add_run("/")
    footer.paragraphs[0]._p.append(parse_xml(docx_field_num_pages))

#Transcrever valor para real por extenso
def valor_por_extenso(valor):
    valor = valor.replace(".", "").replace(",", ".")
    valor_float = float(valor)

    valor_inteiro = int(valor_float)
    valor_centavos = int(round((valor_float - valor_inteiro) * 100))

    if valor_centavos > 0:
        valor_extenso = num2words(valor_inteiro, lang='pt_BR') + ' reais e ' + num2words(valor_centavos, lang='pt_BR') + ' centavos'
    else:
        valor_extenso = num2words(valor_inteiro, lang='pt_BR') + ' reais'

    return valor_extenso

#Transcrever porcentagem para texto
def porcentagem_para_texto(porcentagem):
    # Divide a porcentagem em partes inteiras e decimais
    partes = str(porcentagem).split('.')
    parte_inteira = int(partes[0])
    parte_decimal = int(partes[1])

    # Converte as partes para palavras
    inteira_em_palavras = num2words(parte_inteira, lang='pt_BR').upper()
    decimal_em_palavras = num2words(parte_decimal, lang='pt_BR').upper()

    # Verifica se as partes são singulares ou plurais
    inteira_singular_ou_plural = 'INTEIRO' if parte_inteira == 1 else 'INTEIROS'
    decimal_singular_ou_plural = 'CENTÉSIMO' if parte_decimal == 1 else 'CENTÉSIMOS'

    # Retorna a porcentagem como texto
    return f'{inteira_em_palavras} {inteira_singular_ou_plural} E {decimal_em_palavras} {decimal_singular_ou_plural} POR CENTO'

#Inserir imagem no cabeçalho
def add_imagem_cabecalho(doc):
    caminho_imagem =f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/logo.png'
    # Acesse a seção do documento
    section = doc.sections[0]

    # Crie um cabeçalho
    header = section.header

    # Adicione um parágrafo ao cabeçalho
    paragraph = header.paragraphs[0]

    # Adicione a imagem ao parágrafo
    run = paragraph.add_run()
    run.add_picture(caminho_imagem, width=Inches(1.5))

    # Alinhe o parágrafo à direita
    paragraph.alignment = 2  # 0 (left), 1 (center), 2 (right)
#========================================CRIACAO DOS DOCUMENTOS=====================================================================
def prorroga_data():
    caminho_aditivo = 'C:/Users/cayo_garces/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    prorroga_data = 'prorroga_data.docx'
    # Alteração de vencimento
    doc = Document()
    pCabecalho = doc.add_paragraph()
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)
    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)
        doc.add_paragraph('Clausula Primeira: Pelo presente Aditivo, o EMITENTE e a COOPERATIVA, com a anuência do(s) AVALISTA(S) E CÔNJUGE(S), todos qualificados acima, resolvem de comum acordo as alterações a seguir pactuadas, que passarão a integrar o título ora aditado:').alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pParcelas = doc.add_paragraph()
    novadata = dtPrimParcela
    for i in range(qtdParcelas):
        novadata = novadata + relativedelta(months=+periodicidade)
        data_formatada = novadata.strftime("%d/%m/%Y")
        texto = (str(primParcela+i)+') O vencimento da parcela '+str(i+1)+' foi alterado para '+data_formatada+'\n')
        pParcelas.add_run(texto)        

    txtClausula2='Cláusula Segunda: O vencimento final desta cédula passa a ser '+data_formatada+'.\n'
    txtClausula3='Cláusula Terceira: O valor das parcelas e de eventuais tributos incidentes nesta cédula pode ser alterado devido a mudança na data de vencimento e do recálculo do saldo devedor do título.\n'
    txtClausula4='Cláusula Quarta: O presente aditivo em nada altera as disposições sobre eventuais garantias constituídas sob esta cédula ou outro instrumento, permanecendo válidas e vigentes, nos termos contratados, liberadas somente após a liquidação da operação, conforme o caso.\n'
    txtClausula5='Cláusula Quinta: As demais cláusulas e condições da cédula, não expressamente alteradas neste Aditivo, ficam EXPRESSAMENTE RATIFICADAS, especialmente os encargos ali pactuados.\nAs partes assinam o presente aditivo em vias de igual conteúdo e forma.\n'
    #txtRodape=(strMunEmitent+'-'+strUFEmitente+', '+ str(novadata.day)+' de '+meses[novadata.month]+ ' de '+str(novadata.year))

    pClausulas = doc.add_paragraph(txtClausula2+txtClausula3+txtClausula4+txtClausula5).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)

    add_num_pag(doc)
    add_imagem_cabecalho(doc)
    # Salve o documento
    doc.save(caminho_aditivo+prorroga_data)

def altera_taxa():
    caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    altera_taxa = 'altera_taxa.docx'
    # Alteração de vencimento
    doc = Document()
    pCabecalho = doc.add_paragraph()
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)

    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Pelo presente Aditivo, o ASSOCIADO e a COOPERATIVA, com a anuência do(s) AVALISTA(S) E CÔNJUGE(S), todos qualificados acima, resolvem, de comum acordo e sem o intuito de novar, ')
    sub = pClausulas.add_run('alterar')
    sub.bold, sub.underline = True, True
    pClausulas.add_run(' a cláusula abaixo, que passa a vigorar com a seguinte redação:\n')
    pClausulas.add_run('CLÁUSULA PRIMEIRA').bold=True
    pClausulas.add_run(' – As partes resolvem ')
    pClausulas.add_run('ALTERAR ').bold=True
    pClausulas.add_run('a Cláusula “Forma de Pagamento”, que passará a viger com a redação a seguir:')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Parágrafo Primeiro: O limite e o vencimento desta cédula serão prorrogáveis sucessiva e automaticamente por iguais períodos, a critério da COOPERATIVA, sem a necessidade de qualquer formalidade, o mesmo ocorrendo ao final da primeira e demais prorrogações, permanecendo em vigor todas as cláusulas e condições previstas.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Parágrafo Segundo: Mediante comunicação de uma parte a outra, com 05 (cinco) dias de antecedência, é facultado à COOPERATIVA e/ou ao ASSOCIADO, a qualquer tempo, resilir imotivadamente esta Cédula, encerrando o limite de crédito contratado, ou reduzir o referido limite. Nesses casos, o ASSOCIADO deverá recompor o limite, seja liquidando integralmente o saldo devedor, no caso de encerramento de limite, seja readequando o limite, no caso de redução.”')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('ENCARGOS REMUNERATÓRIOS').bold=True
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('CLÁUSULA SEGUNDA').bold=True
    pClausulas.add_run(' – A partir desta alteração, sobre o saldo médio devedor apresentado ao final de cada mês, incidirão juros apurados pela taxa referencial DI-Cetip Over (Extra-Grupo), calculada e divulgada pela B3 com base nas operações de emissão de depósitos interfinanceiros, ou, no caso de interrupção da sua divulgação, por outra taxa referencial de juros com base equivalente que venha a substituí-la, acrescida de '+str(taxa_ano)+'% ('+porcentagem_para_texto(taxa_ano)+') ao ano, correspondente a '+str(taxa_mes)+'% ('+porcentagem_para_texto(taxa_mes)+')) ao mês, capitalizados mensalmente, no vencimento, nas amortizações, na alteração de vencimento e na liquidação da dívida.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Parágrafo Único: Os encargos remuneratórios apurados serão debitados mensalmente na conta de depósito do(s) EMITENTE(S) onde está cadastrado este limite, no 1. dia útil do mês subsequente ao da apuração, bem como na alteração de vencimento e na liquidação da dívida. Na hipótese de liquidação, alteração de vencimento, ou amortização do empréstimo fora do dia de referência, incidirá atualização "pro rata" dia útil, com utilização da remuneração acumulada do DI-Cetip Over (Extra-Grupo) desde a última atualização, a qual serão somados, proporcionalmente, os encargos denominados adicionais.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('CLÁUSULA TERCEIRA').bold=True
    pClausulas.add_run(' – DA RATIFICAÇÃO: As demais cláusulas e condições da Cédula, não expressamente alteradas neste Aditivo, ficam EXPRESSAMENTE RATIFICADAS, especialmente a(s) garantia(s) constituída(s), que àquela se integra, continuando a produzir seus jurídicos e legais efeitos.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('As Partes ajustam que o presente aditivo e anexos poderão ser assinados digital ou eletronicamente, produzindo todos os efeitos. Nos termos do art. 10, § 2º, da Medida Provisória nº 2.200-2, as Partes expressamente concordam em utilizar e reconhecem como válida qualquer forma de comprovação de anuência aos termos ora acordados em formato eletrônico, ainda que não utilizem de certificado digital emitido no padrão ICP-Brasil, incluindo assinaturas eletrônicas em plataforma específica disponibilizada pelo Sicredi diretamente ou por terceiros. A formalização das avenças na maneira supra acordada será suficiente para a validade e integral vinculação das partes ao presente Contrato.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)
    
    add_imagem_cabecalho(doc)
    add_num_pag(doc)
    # Salve o documento
    doc.save(caminho_aditivo+altera_taxa)

def troca_garantia(add_prorroga_data_flag=False):
    #troca de garantia
    caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    troca_garantia = 'troca_garantia.docx'

    doc = Document()
    pCabecalho = doc.add_paragraph()
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)

    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Pelo presente Aditivo, os presentes signatários, têm, entre si, justas e acordadas as alterações a seguir pactuadas, que passarão a integrar o título ora aditado:')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #INCLUSÃO DA GARANTIA========================================================================================================================================================
    #Imovel
    if strNovaGarantia.lower() in 'imóvel':
        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('CLÁUSULA PRIMEIRA: DA INCLUSÃO DE IMÓVEL EM ALIENAÇÃO FIDUCIÁRIA - As Partes desejam, sem prejuízo da alienação fiduciária de imóvel já constituída, incluir mais um imóvel em alienação fiduciária de imóvel ao Contrato, em garantia do pagamento das dívidas contraídas decorrentes de todas as Operações Financeiras Derivadas, nos termos das Leis 9.514/97 e 13.476/17, aplicando-se a garantia sobre este imóvel todas as cláusulas do Contrato, em especial as aplicáveis à alienação fiduciária de imóvel:')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('1. Descrição do imóvel: (descrição completa conforme matrícula)\n'+nova_garantia)
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('2. Prazo de carência para expedição de intimação: '+strCarencia+' dias')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('3. Valor de avaliação do imóvel para fins de garantia e venda em público leilão: R$ '+strValorImovel+'( '+valor_por_extenso(strValorImovel)+')')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('O EMITENTE e os intervenientes garantidores, caso forem pessoas jurídicas, assim identificados neste documento através do CNPJ, ou quando se tratar de operações de crédito rural, DECLARAM, sob as penas da lei, que o(s) bem(ns) dado(s) em garantia fiduciária por meio deste instrumento NÃO são essenciais à sua atividade empresarial, não se aplicando, dessa forma, as restrições à sua retirada dispostas na parte final do §3º do art. 49 da Lei nº 11.101/05.')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('É responsabilidade do DEVEDOR FIDUCIANTE, efetuar o pagamento de todos os tributos, tarifa(s), despesa(s) e demais encargos relativos ao presente instrumento, além dos impostos que recaiam sobre o(s) bem(ns) dado(s) em garantia de alienação fiduciária, taxa de licenciamento, seguro obrigatório, encargos, multas e demais despesas que incidam ou venham a incidir, direta ou indiretamente, bem como outras despesas judiciais e extrajudiciais decorrentes de depósito, execução ou outra medida que se faça necessária, podendo a COOPERATIVA exigir comprovação, a qualquer tempo. Caso a COOPERATIVA seja demandada a pagar qualquer valor de responsabilidade do DEVEDOR FIDUCIANTE, esta fica desde já autorizada a efetuar o débito dos valores em qualquer conta mantida pelo ASSOCIADO. Em tal hipótese, se não houver saldo disponível, a COOPERATIVA poderá proceder o imediato vencimento antecipado de suas obrigações.')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('CLÁUSULA SEGUNDA: DAS OPERAÇÕES FINANCEIRAS DERIVADAS VINCULADAS AO CONTRATO DE LIMITE DE CRÉDITO - Além dos empréstimos e financiamentos tomados pelo DEVEDOR a partir da assinatura deste instrumento, serão consideradas Operações Financeiras Derivadas firmadas entre o CREDOR e o DEVEDOR e, portanto, abrangidas pelas obrigações e garantias constituídas neste Contrato, os seguintes empréstimos e financiamentos firmados antes da assinatura deste Aditivo ao Contrato pelo CREDOR e pelo DEVEDOR:')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('NÚMERO: '+strTituloExcGarantia)
        pClausulasImovel.add_run('\nINSTRUMENTO: '+strInstrumentoExc)
        pClausulasImovel.add_run('\nVALOR PRINCIPAL: R$ '+strValorExcGarantia)
        pClausulasImovel.add_run('\nVENCIMENTO FINAL: '+strVencFinal)
        pClausulasImovel.add_run('\nTAXA DE JUROS: '+strTxJurosExc+'%')

    #Veiculo
    if strNovaGarantia.lower() in 'veículo':
        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('1. O proprietário abaixo qualificado dá em Alienação Fiduciária, nos termos do artigo 55 da Lei 10.931/04 e do Decreto-Lei 911/69, os bens abaixo discriminados.')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Em caso de veículos automotores, o EMITENTE, deve se dirigir imediatamente ao Centro de Registro de Veículos Automotores - CRVA para a emissão de novos documentos do veículo (CRV/CRLV), sob pena de impossibilitar a posterior baixa do gravame, conforme normas do respectivo DETRAN.')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Proprietário: '+strEmitente+'\nDescrição: '+nova_garantia+'\nLocal de depósito: '+local_deposito+'')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #Avalista
    if strNovaGarantia.lower() in 'avalista':
        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('1 – Por meio deste Aditivo, o EMITENTE e o CREDOR, qualificados acima, decidem, de comum acordo e sem o intuito de novar, incluir à Cédula de Crédito '+strTitulo+' ora aditada o(s) Avalista(s) e respectivo(s) cônjuge(s) abaixo relacionados:')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('Avalista: '+strNomeAvalista)
        pClausulasAvalista.add_run('\nNacionalidade: '+strNacioAvalista)
        pClausulasAvalista.add_run('\nProfissão: '+strProfAvalista)
        pClausulasAvalista.add_run('\nEstado Civil: '+strEndAvalista)
        pClausulasAvalista.add_run('\nRG: '+strRGAvalista)
        pClausulasAvalista.add_run('\nCPF/CNPJ: '+strCPFAvalista)
        pClausulasAvalista.add_run('\nEndereço/CEP: '+strEndAvalista)

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('2 - Em decorrência da inclusão acima, o(s) Avalista(s) ora adicionados, comparecem e assinam o presente Aditivo para assegurar o cumprimento de todas as obrigações principais e acessórias assumida(s) na Cédula ora aditada, na forma da legislação cambial.')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #RETIRADA DA GARANTIA========================================================================================================================================================
    #Imóvel
    if strGarantiaRetirada.lower() in 'imóvel':
        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('CLÁUSULA TERCEIRA: DA ANUÊNCIA CONDICIONADA - O CREDOR compromete-se em autorizar o cancelamento e baixa da alienação fiduciária em garantia do '+strGarantiaRetirada+', constituída por meio de Contrato de Limite de Crédito, nos termos da Lei nº 13.476, de 2017, com pacto adjeto de Alienação Fiduciária de Imóvel para Garantia de Obrigações “em ser” e futuras, datado de '+strDataDocumento+', após, e somente após, o registro da garantia de alienação fiduciária a ser constituída sobre o bem imóvel descrito na Cláusula Primeira acima, no(s) seu(s) registro(s) competente(s)')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #Veiculo
    if strGarantiaRetirada.lower() in 'veículo':
        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('2. Após, e somente após, o(s) registro(s) e/ou averbação da(s) garantia(s) constituída(s) acima, no(s) seu(s) registro(s) competente(s), a ALIENAÇÃO FIDUCIÁRIA do(s) bem(ns) descrito(s) abaixo fica CANCELADA, sendo que este aditivo é o documento hábil para o cancelamento do respectivo registro desta(s) garantia(s):')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Descrição: '+garantia_retirada)
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #Avalista
    if strGarantiaRetirada.lower() in 'avalista':
        if strNovaGarantia.lower() in 'avalista':
            pClausulasAvalista = doc.add_paragraph()
            pClausulasAvalista.add_run('3 – Em decorrência da inclusão do avalista acima, resta cancelado o aval prestado originariamente pela pessoa descrita a seguir:')
            pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY     
        else:
            pClausulasAvalista = doc.add_paragraph()
            pClausulasAvalista.add_run('3 – Após, e somente após, o(s) registro(s) e/ou averbação da(s) garantia(s) constituída(s) acima, resta cancelado o aval prestado originariamente pela pessoa descrita a seguir:')
            pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY 

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('Avalista: '+strNomeExcAvalista)
        pClausulasAvalista.add_run('\nNacionalidade: '+strNacioExcAvalista)
        pClausulasAvalista.add_run('\nProfissão: '+strProfExcAvalista)
        pClausulasAvalista.add_run('\nEstado Civil: '+strEndExcAvalista)
        pClausulasAvalista.add_run('\nRG: '+strRGExcAvalista)
        pClausulasAvalista.add_run('\nCPF/CNPJ: '+strCPFExcAvalista)
        pClausulasAvalista.add_run('\nEndereço/CEP: '+strEndExcAvalista)

    pClausulas = doc.add_paragraph()
    #caso possuir prorrogação de data
    if add_prorroga_data_flag:
        add_prorroga_data()
    pClausulas.add_run('Os signatários ratificam o título em todos os seus termos, itens e condições não expressamente alterados por este documento, que àquela se integra, formando um todo único e indivisível para todos os efeitos.As partes assinam o presente Aditivo em 2 (dois) vias de igual conteúdo e forma.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)
    add_imagem_cabecalho(doc)
    add_num_pag(doc)
    # Salve o documento
    doc.save(caminho_aditivo+troca_garantia)

def exclui_garantia(add_prorroga_data_flag=False):
    #troca de garantia veiculos
    caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    troca_garantia = 'exclui_garantia.docx'

    doc = Document()
    pCabecalho = doc.add_paragraph()
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)

    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Pelo presente Aditivo, os presentes signatários, têm, entre si, justas e acordadas as alterações a seguir pactuadas, que passarão a integrar o título ora aditado:')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    #Adaptação caso seja veiculo ou avalista
    if strGarantiaRetirada.lower() in 'imóvel':
        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('III – DAS CLÁUSULAS')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('1º) Com base no Inciso IV do artigo 1.499 do Código Civil, a hipoteca cedular do '+strGarantiaRetirada+' é cancelada, sendo que este aditivo é o documento hábil para o cancelamento do gravame registrado na matrícula do imóvel.')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #caso possuir prorrogação de data
        if add_prorroga_data_flag:
            add_prorroga_data()

    if strGarantiaRetirada.lower() in 'veículo':
        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('1. A ALIENAÇÃO FIDUCIÁRIA do(s) bem(s) descrito(s) abaixo é CANCELADA, sendo que este aditivo é o documento hábil para o cancelamento do respectivo registro desta garantia.')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Descrição: '+garantia_retirada)
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('2. Após, e somente após, o(s) registro(s) e/ou averbação da(s) garantia(s) constituída(s) acima, no(s) seu(s) registro(s) competente(s), a ALIENAÇÃO FIDUCIÁRIA do(s) bem(ns) descrito(s) abaixo fica CANCELADA, sendo que este aditivo é o documento hábil para o cancelamento do respectivo registro desta(s) garantia(s):')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Descrição: '+garantia_retirada)
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #caso possuir prorrogação de data
        if add_prorroga_data_flag:
            add_prorroga_data()
    if strGarantiaRetirada.lower() in 'avalista':
        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('CLÁUSULA PRIMEIRA – DO CANCELAMENTO DO AVAL As Partes, qualificadas acima, resolvem, de comum acordo e sem o intuito de novar, promover o CANCELAMENTO DO AVAL prestado originalmente pela pessoa descrita a seguir:')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY     

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('Avalista: '+strNomeExcAvalista)
        pClausulasAvalista.add_run('\nNacionalidade: '+strNacioExcAvalista)
        pClausulasAvalista.add_run('\nProfissão: '+strProfExcAvalista)
        pClausulasAvalista.add_run('\nEstado Civil: '+strEndExcAvalista)
        pClausulasAvalista.add_run('\nRG: '+strRGExcAvalista)
        pClausulasAvalista.add_run('\nCPF/CNPJ: '+strCPFExcAvalista)
        pClausulasAvalista.add_run('\nEndereço/CEP: '+strEndExcAvalista)
        if json_excavalista[0]['vRegCas']!=None:
            pClausulasAvalista.add_run('Cônjuge do Avalista: '+json_excavalista[0]['vNomeConjuge'])
            pClausulasAvalista.add_run('\nNacionalidade: '+json_excavalista[0]['vNacConj'])
            pClausulasAvalista.add_run('\nProfissão: '+json_excavalista[0]['vProfConj'])
            pClausulasAvalista.add_run('\nEstado Civil: '+json_excavalista[0]['vEstCivil'])
            pClausulasAvalista.add_run('\nRG: '+json_excavalista[0]['vRGConj'])
            pClausulasAvalista.add_run('\nCPF/CNPJ: '+json_excavalista[0]['vCPFConjuge'])
            pClausulasAvalista.add_run('\nEndereço/CEP: '+json_excavalista[0]['vEndConj']+', '+json_excavalista[0]['vNumConj']+', '+json_excavalista[0]['vBairroConj']+', '+json_excavalista[0]['vCEPConj']+', '+json_excavalista[0]['vCidConj']+', '+json_excavalista[0]['vUFConj'])
        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('CLÁUSULA SEGUNDA – DA RATIFICAÇÃO')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #caso possuir prorrogação de data
        if add_prorroga_data_flag:
            add_prorroga_data()
        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('As demais cláusulas e condições da Cédula, não expressamente alteradas neste Aditivo, ficam expressamente ratificadas, especialmente a(s) garantia(s) constituída(s) originalmente, que àquela se integra, continuando a produzir seus jurídicos e legais efeitos.E por estarem de inteiro e comum acordo com as cláusulas e condições deste Aditivo, as Partes assinam o presente aditivo e anexos que poderão ser assinados digital ou eletronicamente, produzindo todos os efeitos. Nos termos do art. 10, § 2º, da Medida Provisória nº 2.200-2, as Partes expressamente concordam em utilizar e reconhecem como válida qualquer forma de comprovação de anuência aos termos ora acordados em formato eletrônico, ainda que não utilizem de certificado digital emitido no padrão ICP-Brasil, incluindo assinaturas eletrônicas em plataforma específica disponibilizada pelo Sicredi diretamente ou por terceiros. A formalização das avenças na maneira supra acordada será suficiente para a validade e integral vinculação das partes ao presente Contrato.')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Os signatários ratificam o título em todos os seus termos, itens e condições não expressamente alterados por este documento, que àquela se integra, formando um todo único e indivisível para todos os efeitos.As partes assinam o presente Aditivo em 2 (dois) vias de igual conteúdo e forma.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)
    add_imagem_cabecalho(doc)
    add_num_pag(doc)
    # Salve o documento
    doc.save(caminho_aditivo+troca_garantia)

def inclui_garantia(add_prorroga_data_flag=False):
    #troca de garantia veiculos
    caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    troca_garantia = 'inclui_garantia.docx'

    doc = Document()
    pCabecalho = doc.add_paragraph()
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)

    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Pelo presente Aditivo, os presentes signatários, têm, entre si, justas e acordadas as alterações a seguir pactuadas, que passarão a integrar o título ora aditado:')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    #Adaptação caso seja veiculo ou avalista
    if strNovaGarantia.lower() in 'imóvel':
        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('III – DAS CLÁUSULAS')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('1º) Esta operação de crédito é uma Operação Financeira Derivada do Contrato de Limite de Crédito firmado entre o CREDOR e o ASSOCIADO no dia 01/06/2022, nos termos da Lei nº 13.476, de 2017, e, portanto, as obrigações principais e acessórias assumidas neste instrumento estão garantidas pela alienação fiduciária do imóvel registrado conforme a seguir:')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run(strNovaGarantia)
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #caso possuir prorrogação de data
        if add_prorroga_data_flag:
            add_prorroga_data()
        pClausulasImovel = doc.add_paragraph()
        pClausulasImovel.add_run('O EMITENTE e os intervenientes garantidores, caso forem pessoas jurídicas, assim identificados neste documentoatravés do CNPJ, ou quando se tratar de operações de crédito rural, DECLARAM, sob as penas da lei, que o(s) bem(ns) dado(s) em garantia fiduciária por meio deste instrumento NÃO são essenciais à sua atividade empresarial, não se aplicando, dessa forma, as restrições à sua retirada dispostas na parte final do §3º do art. 49 da Lei nº 11.101/05.')
        pClausulasImovel.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if strNovaGarantia.lower() in 'veículo':
        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('1º) O proprietário abaixo qualificado dá em Alienação Fiduciária, nos termos do artigo 55 da Lei 10.931/04 e do Decreto-Lei 911/69, os bens abaixo discriminados:')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Descrição: '+strNovaGarantia)
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('Em caso de veículos automotores, o EMITENTE, deve se dirigir imediatamente ao Centro de Registro de Veículos Automotores - CRVA para a emissão de novos documentos do veículo (CRV/CRLV), sob pena de impossibilitar a posterior baixa do gravame, conforme normas do respectivo DETRAN.')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #caso possuir prorrogação de data
        if add_prorroga_data_flag:
            add_prorroga_data()
        pClausulasVeiculo = doc.add_paragraph()
        pClausulasVeiculo.add_run('A inclusão da garantia efetuada acima é constituída em conjunto com as demais garantias eventualmente constituídas na Cédula, sem substituí-las. Os signatários ratificam o título em todos os seus termos, itens e condições não expressamente alterados por este documento, que àquela se integra, formando um todo único e indivisível para todos os efeitos.')
        pClausulasVeiculo.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    if strNovaGarantia.lower() in 'avalista':
        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('III – DAS CLÁUSULAS')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY     

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('1 – Por meio deste Aditivo, o EMITENTE e o CREDOR, qualificados acima, decidem, de comum acordo e sem o intuito de novar, incluir à Cédula de Crédito '+strTitulo+' ora aditada o(s) Avalista(s) e respectivo(s) cônjuge(s) abaixo relacionados:')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY    

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('Avalista: '+strNomeIncAvalista)
        pClausulasAvalista.add_run('\nNacionalidade: '+strNacioIncAvalista)
        pClausulasAvalista.add_run('\nProfissão: '+strProfIncAvalista)
        pClausulasAvalista.add_run('\nEstado Civil: '+strEndIncAvalista)
        pClausulasAvalista.add_run('\nRG: '+strRGIncAvalista)
        pClausulasAvalista.add_run('\nCPF/CNPJ: '+strCPFIncAvalista)
        pClausulasAvalista.add_run('\nEndereço/CEP: '+strEndIncAvalista)

        if json_incavalista[0]['vRegCas']!=None:
            pClausulasAvalista.add_run('Cônjuge do Avalista: '+json_incavalista[0]['vNomeConjuge'])
            pClausulasAvalista.add_run('\nNacionalidade: '+json_incavalista[0]['vNacConj'])
            pClausulasAvalista.add_run('\nProfissão: '+json_incavalista[0]['vProfConj'])
            pClausulasAvalista.add_run('\nEstado Civil: '+json_incavalista[0]['vEstCivil'])
            pClausulasAvalista.add_run('\nRG: '+json_incavalista[0]['vRGConj'])
            pClausulasAvalista.add_run('\nCPF/CNPJ: '+json_incavalista[0]['vCPFConjuge'])
            pClausulasAvalista.add_run('\nEndereço/CEP: '+json_incavalista[0]['vEndConj']+', '+json_incavalista[0]['vNumConj']+', '+json_incavalista[0]['vBairroConj']+', '+json_incavalista[0]['vCEPConj']+', '+json_incavalista[0]['vCidConj']+', '+json_incavalista[0]['vUFConj'])

        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('2 - Em decorrência da inclusão acima, o(s) Avalista(s) ora adicionados, comparecem e assinam o presente Aditivo para assegurar o cumprimento de todas as obrigações principais e acessórias assumida(s) na Cédula ora aditada, na forma da legislação cambial.')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #caso possuir prorrogação de data
        if add_prorroga_data_flag:
            add_prorroga_data()
        pClausulasAvalista = doc.add_paragraph()
        pClausulasAvalista.add_run('3 - As demais cláusulas e condições da Cédula, não expressamente alteradas neste Aditivo, ficam expressamente ratificadas, especialmente a(s) garantia(s) constituída(s) originalmente, que àquela se integra, continuando a produzir seus jurídicos e legais efeitos.E por estarem de inteiro e comum acordo com as cláusulas e condições deste Aditivo, as Partes assinam o presente aditivo e anexos que poderão ser assinados digital ou eletronicamente, produzindo todos os efeitos. Nos termos do art. 10, § 2º, da Medida Provisória nº 2.200-2, as Partes expressamente concordam em utilizar e reconhecem como válida qualquer forma de comprovação de anuência aos termos ora acordados em formato eletrônico, ainda que não utilizem de certificado digital emitido no padrão ICP-Brasil, incluindo assinaturas eletrônicas em plataforma específica disponibilizada pelo Sicredi diretamente ou por terceiros. A formalização das avenças na maneira supra acordada será suficiente para a validade e integral vinculação das partes ao presente Contrato.')
        pClausulasAvalista.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Os signatários ratificam o título em todos os seus termos, itens e condições não expressamente alterados por este documento, que àquela se integra, formando um todo único e indivisível para todos os efeitos.As partes assinam o presente Aditivo em 2 (dois) vias de igual conteúdo e forma.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)
    add_imagem_cabecalho(doc)
    add_num_pag(doc)
    # Salve o documento
    doc.save(caminho_aditivo+troca_garantia)

def altera_vinculo():
    #altera vinculo
    caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    altera_vinculo = 'altera_vinculo.docx'

    doc = Document()
    pCabecalho = doc.add_paragraph()
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)

    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)
    #
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('DADOS DO INSTRUMENTO DE CRÉDITO: ')
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('NÚMERO DO TÍTULO: '+strTitulo)
    pClausulas.add_run('\nDATA EMISSÃO: '+strdataemissao)
    
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Pelo presente Aditivo, as Partes qualificadas acima, têm, entre si, justas e acordadas as alterações a seguir pactuadas, que passarão a integrar o título ora aditado:')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('CLÁUSULA PRIMEIRA - As Partes estabelecem que o saldo devedor, indicado acima, será pago em '+strQtdParcelas+' parcelas, iguais e sucessivas, de R$ '+str(strValorParcela)+' ('+valor_por_extenso(strValorParcela)+') , parcelas estas que incluem o principal e os encargos calculados sobre o saldo devedor, cada uma conforme o cronograma a seguir: '+strdataparcela+'.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run(f"CLÁUSULA SEGUNDA - Sobre o saldo devedor acima, incidem juros à taxa efetiva de {strTaxaAno}% () ao ano (2,300000% ao mês), capitalizados mensalmente, calculados de acordo com a Tabela PRICE, e já adicionados nas parcelas indicadas acima.")
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('CLÁUSULA TERCEIRA - Com o fim único e exclusivo de ajuste nos registros e controles do CREDOR, sem constituir qualquer novação, baixa, cancelamento ou extinção das obrigações estabelecidas na Cédula e de suas garantias, o número da Cédula, constante em seu cabeçalho, é alterado neste ato, passando a ser o nº '+strNovoTitulo+'.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()    
    pClausulas.add_run('CLÁUSULA QUARTA - Os signatários ratificam a Cédula e os seus registros em todos os seus termos, itens e condições não expressamente alterados por este documento, especialmente da(s) garantia(s) constituída(s), que àquela se integra, formando um todo único e indivisível para todos os efeitos.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()    
    pClausulas.add_run(nova_garantia)
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()    
    pClausulas.add_run('As partes assinam o presente Aditivo em 2(DUAS) vias de igual conteúdo e forma.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #
    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)
    add_imagem_cabecalho(doc)
    add_num_pag(doc)
    # Salve o documento
    doc.save(caminho_aditivo+altera_vinculo)

def renova_vencimento():
    #altera vinculo
    caminho_aditivo = f'C:/Users/{username}/Sicredi/01.0821 Sede Univales - Documentos/General/sidi_relatorios/rpa_0821_aditivo/'
    renova_vencimento = 'renova_vencimento.docx'

    doc = Document()
    pCabecalho = doc.add_paragraph()
    txtCabecalho = 'Aditivo à Cédula de Crédito Bancário abertura de limite de crédito rotativo nº '+strTitulo+', emitida por '+strEmitente+' em favor da COOPERATIVA DE CRÉDITO POUPANCA E INVESTIMENTO UNIVALES SICREDI UNIVALES MT/RO, CNPJ 70.431.630/0001-04, em 21 de setembro de 2022.'
    pCabecalho.add_run(txtCabecalho).bold = True
    pCabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pCredor = doc.add_paragraph()
    pCredor.add_run(txtCredor)
    pCredor.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pEmitente = doc.add_paragraph()
    pEmitente.add_run(txtEmitente)
    pEmitente.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_interveniente_garantidor(doc,json_intergarantidor)

    #SE HOUVER AVALISTA
    if json_avalista:
        add_avalistas(doc, json_avalista)
    #
    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Pelo presente Aditivo, o ASSOCIADO e a COOPERATIVA, com a anuência do(s) AVALISTA(S) E CÔNJUGE(S), todos qualificados acima, resolvem, de comum acordo e sem o intuito de novar, alterar a cláusula abaixo, que passa a vigorar com a seguinte redação:')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('“FORMA DE PAGAMENTO – O limite de crédito ora liberado deverá ser resgatado, integralmente, pelos valores então utilizados, em 08/09/2024, sendo, durante a vigência desta cédula, recomposto automaticamente na medida em que forem efetuados pagamentos.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Parágrafo Primeiro: O limite e o vencimento desta cédula serão prorrogáveis sucessiva e automaticamente por iguais períodos, a critério da COOPERATIVA, sem a necessidade de qualquer formalidade, o mesmo ocorrendo ao final da primeira e demais prorrogações, permanecendo em vigor todas as cláusulas e condições previstas.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas.paragraph_format.left_indent = Inches(0.5)

    pClausulas = doc.add_paragraph()
    pClausulas.add_run('Parágrafo Segundo: Mediante comunicação de uma parte a outra, com 05 (cinco) dias de antecedência, é facultado à COOPERATIVA e/ou ao ASSOCIADO, a qualquer tempo, resilir imotivadamente esta Cédula, encerrando o limite de crédito contratado, ou reduzir o referido limite. Nesses casos, o ASSOCIADO deverá recompor o limite, seja liquidando integralmente o saldo devedor, no caso de encerramento de limite, seja readequando o limite, no caso de redução.”')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas.paragraph_format.left_indent = Inches(0.5)

    pClausulas = doc.add_paragraph()    
    pClausulas.add_run('Parágrafo Único: Os encargos remuneratórios apurados serão debitados mensalmente na conta de depósito do(s) EMITENTE(S) onde está cadastrado este limite, no 1. dia útil do mês subsequente ao da apuração, bem como na alteração de vencimento e na liquidação da dívida. Na hipótese de liquidação, alteração de vencimento, ou amortização do empréstimo fora do dia de referência, incidirá atualização "pro rata" dia útil, com utilização da remuneração acumulada do DI-Cetip Over (Extra-Grupo) desde a última atualização, a qual serão somados, proporcionalmente, os encargos denominados adicionais.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()    
    pClausulas.add_run('CLÁUSULA PRIMEIRA – DA RATIFICAÇÃO: As demais cláusulas e condições da Cédula, não expressamente alteradas neste Aditivo, ficam EXPRESSAMENTE RATIFICADAS, especialmente a(s) garantia(s) constituída(s), que àquela se integra, continuando a produzir seus jurídicos e legais efeitos.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    pClausulas = doc.add_paragraph()    
    pClausulas.add_run(nova_garantia)
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    pClausulas = doc.add_paragraph()    
    pClausulas.add_run('As Partes ajustam que o presente aditivo e anexos poderão ser assinados digital ou eletronicamente, produzindo todos os efeitos. Nos termos do art. 10, § 2º, da Medida Provisória nº 2.200-2, as Partes expressamente concordam em utilizar e reconhecem como válida qualquer forma de comprovação de anuência aos termos ora acordados em formato eletrônico, ainda que não utilizem de certificado digital emitido no padrão ICP-Brasil, incluindo assinaturas eletrônicas em plataforma específica disponibilizada pelo Sicredi diretamente ou por terceiros. A formalização das avenças na maneira supra acordada será suficiente para a validade e integral vinculação das partes ao presente Contrato.')
    pClausulas.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    #
    pRodape = doc.add_paragraph(txtRodape).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    pAssinaturas = doc.add_paragraph()
    pAssinaturas.add_run(txtAssCredor)

    pAssinaturas.add_run(txtAssEmitente)

    if json_avalista:
        add_ass_avalistas(doc, json_avalista)
    #SE HOUVER INTERVENIENTE GARANTIDOR
    if json_intergarantidor:
        add_ass_inter_garantidor(doc,json_intergarantidor)
    add_imagem_cabecalho(doc)
    add_num_pag(doc)
    # Salve o documento
    doc.save(caminho_aditivo+renova_vencimento)    

troca_garantia()
#prorroga_data()
#altera_taxa()
#altera_vinculo()
#renova_vencimento()